"""Xuất báo cáo Word (.docx) cho luồng Conversational BI."""

from __future__ import annotations

import io
import logging
import os
import re
import tempfile
from datetime import datetime
from typing import Any

import pandas as pd
import plotly.io as pio
from docx import Document

_logger = logging.getLogger(__name__)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _set_cell_border(cell: Any, **kwargs: str) -> None:
    """Gắn border đơn giản cho ô bảng Word."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right"):
        edge_val = kwargs.get(edge)
        if edge_val is None:
            continue
        tag = f"w:{edge}"
        element = OxmlElement(tag)
        element.set(qn("w:val"), edge_val)
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        tc_pr.append(element)


def _apply_table_borders(table: Any) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top="single",
                bottom="single",
                left="single",
                right="single",
            )


def _clean_insight_text(text: str) -> str:
    """Bỏ markdown cơ bản trước khi ghi Word."""
    cleaned = (text or "").strip()
    if not cleaned:
        return "Không có nhận định từ AI."
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    return cleaned


def _cell_value(value: Any) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return f"{int(value):,}"
        return f"{value:,.2f}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def is_kaleido_available() -> bool:
    """Kiểm tra kaleido đã cài — cần cho export PNG từ Plotly."""
    try:
        import kaleido  # noqa: F401
        return True
    except ImportError:
        return False


def _figure_to_png_bytes(figure: Any) -> bytes | None:
    """
    Export Plotly Figure → PNG bytes.
    Thử nhiều cách (to_image / write_image) vì môi trường Streamlit/OS khác nhau.
    """
    if figure is None:
        return None

    export_kwargs = {"format": "png", "width": 1200, "height": 675, "scale": 2}

    # Cách 1: fig.to_image với engine kaleido (Plotly 5.x–6.x)
    for engine in ("kaleido", None):
        try:
            kwargs = dict(export_kwargs)
            if engine:
                kwargs["engine"] = engine
            return figure.to_image(**kwargs)
        except Exception as exc:
            _logger.debug("to_image engine=%s failed: %s", engine, exc)

    # Cách 2: plotly.io.write_image qua file tạm (ổn định hơn trên Windows/Streamlit)
    tmp_path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
        pio.write_image(figure, tmp_path, **export_kwargs, engine="kaleido")
        with open(tmp_path, "rb") as img_file:
            return img_file.read()
    except Exception as exc:
        _logger.warning("write_image kaleido failed: %s", exc)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    # Cách 3: write_image không chỉ định engine
    tmp_path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp_path = tmp.name
        pio.write_image(figure, tmp_path, **export_kwargs)
        with open(tmp_path, "rb") as img_file:
            return img_file.read()
    except Exception as exc:
        _logger.warning("write_image default engine failed: %s", exc)
        return None
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def _embed_chart_image(doc: Document, figure: Any) -> bool:
    """
    Chèn ảnh PNG từ Plotly vào Word (kaleido).
    Trả False nếu không export được — Word vẫn tạo bình thường.
    """
    img_bytes = _figure_to_png_bytes(figure)
    if not img_bytes:
        return False
    return _embed_png_bytes(doc, img_bytes)


def _embed_png_bytes(doc: Document, img_bytes: bytes, width_inches: float = 6.0) -> bool:
    try:
        stream = io.BytesIO(img_bytes)
        doc.add_picture(stream, width=Inches(width_inches))
        return True
    except Exception as exc:
        _logger.warning("doc.add_picture failed: %s", exc)
        return False


def embed_base64_png(doc: Document, data_url_or_b64: str | None) -> bool:
    """Nhúng PNG từ data URL (data:image/png;base64,...) hoặc raw base64."""
    if not data_url_or_b64:
        return False
    raw = data_url_or_b64.strip()
    if "," in raw and raw.lower().startswith("data:"):
        raw = raw.split(",", 1)[1]
    try:
        import base64

        img_bytes = base64.b64decode(raw)
    except Exception as exc:
        _logger.warning("base64 decode chart failed: %s", exc)
        return False
    return _embed_png_bytes(doc, img_bytes)


def create_word_report_api(
    query: str,
    insight_text: str,
    dataframe: pd.DataFrame | None,
    chart_image_base64: str | None = None,
    article_markdown: str = "",
) -> tuple[bytes, bool]:
    """
    Export Word cho Next.js API — ưu tiên ảnh ECharts (base64),
    nếu có article_markdown thì layout bài báo.
    """
    if (article_markdown or "").strip():
        doc = Document()
        chart_embedded = False
        parts = split_article_markdown(article_markdown)
        title_text = parts["title"] or "Bài báo phân tích"
        heading = doc.add_heading(title_text, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if parts["lead"]:
            p = doc.add_paragraph(parts["lead"])
            for run in p.runs:
                run.italic = True
        if chart_image_base64:
            chart_embedded = embed_base64_png(doc, chart_image_base64)
            if chart_embedded:
                cap = doc.add_paragraph("Biểu đồ phân tích")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                note = doc.add_paragraph(
                    "Số liệu chi tiết xem tại bảng bên dưới."
                )
                note.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in note.runs:
                    run.italic = True
                    run.font.size = Pt(9)
        for line in _clean_insight_text(parts["body"] or article_markdown).splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
        if dataframe is not None and not dataframe.empty:
            table_title = (
                "Bảng số liệu chi tiết (tương ứng biểu đồ trên)"
                if chart_embedded
                else "Bảng số liệu"
            )
            doc.add_heading(table_title, level=2)
            _add_dataframe_table(doc, dataframe)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue(), chart_embedded

    # Báo cáo insight thường
    doc = Document()
    chart_embedded = False
    title = doc.add_heading("BÁO CÁO PHÂN TÍCH DỮ LIỆU", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    time_p = doc.add_paragraph()
    time_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_p.add_run(f"Thời gian tạo báo cáo: {now}")
    time_run.font.size = Pt(10)
    time_run.italic = True
    doc.add_paragraph()
    doc.add_paragraph(f"Yêu cầu phân tích: {(query or '').strip() or '—'}")
    doc.add_heading("Phần 1 — Executive Summary", level=2)
    for line in _clean_insight_text(insight_text).splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    if chart_image_base64:
        doc.add_heading("Phần 2 — Biểu đồ trực quan", level=2)
        chart_embedded = embed_base64_png(doc, chart_image_base64)
        if chart_embedded:
            note = doc.add_paragraph(
                "Số liệu chi tiết xem tại bảng bên dưới."
            )
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in note.runs:
                run.italic = True
                run.font.size = Pt(9)
        table_heading = "Phần 3 — Số liệu chi tiết (tương ứng biểu đồ trên)"
    else:
        table_heading = "Phần 2 — Bảng số liệu"
    doc.add_heading(table_heading, level=2)
    if dataframe is None or dataframe.empty:
        doc.add_paragraph("Không có dữ liệu bảng để hiển thị.")
    else:
        _add_dataframe_table(doc, dataframe)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), chart_embedded


def split_article_markdown(markdown: str) -> dict[str, str]:
    """
    Tách bài markdown thành title / lead / body.

    - title: dòng # đầu tiên (không gồm dấu #)
    - lead: đoạn văn trước ## đầu tiên (sau title)
    - body: từ ## đầu tiên đến hết
    """
    text = (markdown or "").strip()
    if not text:
        return {"title": "", "lead": "", "body": ""}

    lines = text.splitlines()
    title = ""
    start_idx = 0
    if lines and lines[0].lstrip().startswith("# ") and not lines[0].lstrip().startswith("##"):
        title = lines[0].lstrip()[2:].strip()
        start_idx = 1
        while start_idx < len(lines) and not lines[start_idx].strip():
            start_idx += 1

    rest = "\n".join(lines[start_idx:]).strip()
    # Tìm ## heading đầu tiên
    body_match = re.search(r"(?m)^##\s+", rest)
    if body_match:
        lead = rest[: body_match.start()].strip()
        body = rest[body_match.start() :].strip()
    else:
        # Không có ## — lấy 1–2 đoạn đầu làm lead, phần còn lại body
        paras = re.split(r"\n\s*\n", rest, maxsplit=1)
        lead = (paras[0] or "").strip()
        body = (paras[1] if len(paras) > 1 else "").strip()

    return {"title": title, "lead": lead, "body": body}


def _add_markdownish_paragraphs(doc: Document, text: str) -> None:
    """Ghi text (đã bỏ **bold**) thành paragraph; ## → heading level 2."""
    cleaned = _clean_insight_text(text)
    for line in cleaned.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        else:
            # Bỏ prefix # thừa nếu còn
            if stripped.startswith("#"):
                stripped = stripped.lstrip("#").strip()
            if stripped:
                doc.add_paragraph(stripped)


def _add_dataframe_table(doc: Document, dataframe: pd.DataFrame) -> None:
    safe_df = dataframe.fillna("")
    rows, cols = safe_df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.style = "Table Grid"

    for j, col_name in enumerate(safe_df.columns):
        table.rows[0].cells[j].text = str(col_name)

    for i in range(rows):
        for j in range(cols):
            table.rows[i + 1].cells[j].text = _cell_value(safe_df.iat[i, j])

    _apply_table_borders(table)


def create_word_report(
    query: str,
    insight_text: str,
    dataframe: pd.DataFrame | None,
    figure: Any | None = None,
) -> tuple[bytes, bool]:
    """
    Tạo file Word báo cáo phân tích (Executive Summary + chart + bảng).

    Args:
        query: Câu hỏi gốc.
        insight_text: Nội dung báo cáo AI.
        dataframe: Bảng số liệu.
        figure: Đối tượng Plotly Figure — nhúng PNG nếu kaleido khả dụng.

    Returns:
        (nội dung file .docx dạng bytes, chart_embedded: True nếu có ảnh biểu đồ).
    """
    doc = Document()
    chart_embedded = False

    title = doc.add_heading("BÁO CÁO PHÂN TÍCH DỮ LIỆU", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    time_p = doc.add_paragraph()
    time_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_p.add_run(f"Thời gian tạo báo cáo: {now}")
    time_run.font.size = Pt(10)
    time_run.italic = True

    doc.add_paragraph()
    doc.add_paragraph(f"Yêu cầu phân tích: {(query or '').strip() or '—'}")

    doc.add_heading("Phần 1 — Executive Summary", level=2)
    for line in _clean_insight_text(insight_text).splitlines():
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)

    # Biểu đồ ngay dưới Executive Summary (trước bảng số liệu)
    if figure is not None:
        doc.add_heading("Phần 2 — Biểu đồ trực quan", level=2)
        chart_embedded = _embed_chart_image(doc, figure)
        if not chart_embedded:
            hint = (
                "pip install -U kaleido plotly"
                if not is_kaleido_available()
                else "kiểm tra log server hoặc thử đổi loại biểu đồ (Bar/Line)"
            )
            doc.add_paragraph(
                f"(Không thể nhúng biểu đồ vào file — {hint})"
            )
        table_heading = "Phần 3 — Bảng số liệu"
    else:
        table_heading = "Phần 2 — Bảng số liệu"

    doc.add_heading(table_heading, level=2)

    if dataframe is None or dataframe.empty:
        doc.add_paragraph("Không có dữ liệu bảng để hiển thị.")
    else:
        _add_dataframe_table(doc, dataframe)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), chart_embedded


def create_article_word(
    *,
    query: str,
    article_markdown: str,
    dataframe: pd.DataFrame | None,
    figure: Any | None = None,
    outline: dict[str, Any] | None = None,
    chart_caption: str = "Biểu đồ phân tích",
) -> tuple[bytes, bool]:
    """
    Word bài báo: tiêu đề → lead → ảnh chart → body → bảng số liệu.

    Layout gần Vietstock hơn (ảnh xen sau tóm tắt luận điểm).
    """
    parts = split_article_markdown(article_markdown)
    outline = outline or {}
    title_text = (
        parts["title"]
        or str(outline.get("title") or "").strip()
        or "Bài báo phân tích"
    )

    doc = Document()
    chart_embedded = False

    heading = doc.add_heading(title_text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Thời gian: {now}")
    meta_run.font.size = Pt(10)
    meta_run.italic = True

    if outline.get("angle"):
        angle_p = doc.add_paragraph()
        angle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        angle_run = angle_p.add_run(str(outline["angle"]))
        angle_run.italic = True
        angle_run.font.size = Pt(11)

    style = str(outline.get("style") or "")
    if style == "vietstock":
        badge = doc.add_paragraph()
        badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = badge.add_run("Phong cách: Báo cáo phân tích (Vietstock-inspired)")
        br.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph(f"Yêu cầu phân tích: {(query or '').strip() or '—'}")

    # Lead / tóm tắt luận điểm
    lead = parts["lead"]
    if lead:
        doc.add_heading("Tóm tắt luận điểm", level=2)
        _add_markdownish_paragraphs(doc, lead)
    elif not parts["body"]:
        # Toàn bài không tách được — ghi full trước chart
        doc.add_heading("Nội dung", level=2)
        _add_markdownish_paragraphs(doc, article_markdown)

    # Ảnh biểu đồ ngay sau lead
    if figure is not None:
        doc.add_heading("Biểu đồ", level=2)
        chart_embedded = _embed_chart_image(doc, figure)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(
            chart_caption
            if chart_embedded
            else (
                "(Không nhúng được biểu đồ — "
                + (
                    "pip install -U kaleido plotly"
                    if not is_kaleido_available()
                    else "thử đổi Bar/Line hoặc kiểm tra kaleido"
                )
                + ")"
            )
        )
        cap_run.italic = True
        cap_run.font.size = Pt(9)

    # Body còn lại
    if parts["body"]:
        _add_markdownish_paragraphs(doc, parts["body"])

    # Bảng số liệu cuối
    doc.add_heading("Bảng số liệu", level=2)
    if dataframe is None or dataframe.empty:
        doc.add_paragraph("Không có dữ liệu bảng để hiển thị.")
    else:
        _add_dataframe_table(doc, dataframe)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), chart_embedded
